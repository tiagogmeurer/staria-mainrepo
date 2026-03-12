import os
import re
import hashlib
from pathlib import Path
from typing import Iterable

from tools.embeddings import embed_texts
from rag.vectorstore import upsert, delete_by_doc

# ====== Paths ======
ROOT = Path(os.getenv("DRIVE_SYNC_ROOT", "")).resolve()
if not str(ROOT) or str(ROOT) == ".":
    raise SystemExit(
        'DRIVE_SYNC_ROOT não definido. Rode: setx DRIVE_SYNC_ROOT "G:\\...\\StarIA"'
    )

# Suporta _inbox ou 00_inbox
INBOX = ROOT / "curriculos"

# ====== File types ======
EXTS = {".txt", ".md", ".pdf", ".docx", ".xlsx"}

# ====== Chunking ======
CHUNK_CHARS = 2200
OVERLAP_CHARS = 280
MAX_DOC_CHARS = 200_000  # proteção: não indexar docs gigantescos sem controle


# ====== Helpers ======
def file_hash(p: Path) -> str:
    h = hashlib.sha1()
    with p.open("rb") as f:
        while True:
            b = f.read(1024 * 1024)
            if not b:
                break
            h.update(b)
    return h.hexdigest()


def clean_text(s: str) -> str:
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = re.sub(r"[ \t]+\n", "\n", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def chunk_text(text: str) -> list[str]:
    text = clean_text(text)
    if not text:
        return []

    # corta para evitar explosão em arquivos enormes
    if len(text) > MAX_DOC_CHARS:
        text = text[:MAX_DOC_CHARS] + "\n\n[TRUNCATED]"

    chunks = []
    i = 0
    step = max(1, CHUNK_CHARS - OVERLAP_CHARS)
    while i < len(text):
        chunk = text[i : i + CHUNK_CHARS]
        chunk = chunk.strip()
        if chunk:
            chunks.append(chunk)
        i += step
    return chunks


# ====== Readers ======
def read_txt_md(p: Path) -> str:
    return p.read_text(encoding="utf-8", errors="ignore")


def read_pdf(p: Path) -> str:
    # PDFs digitais ok. PDF escaneado vai vir vazio (OCR depois se precisar).
    from pypdf import PdfReader

    reader = PdfReader(str(p))
    parts: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            t = page.extract_text() or ""
        except Exception:
            t = ""
        t = clean_text(t)
        if t:
            parts.append(f"[PAGE {i+1}]\n{t}")
    return "\n\n".join(parts)


def read_docx(p: Path) -> str:
    import docx  # python-docx

    d = docx.Document(str(p))
    parts: list[str] = []
    for para in d.paragraphs:
        txt = para.text.strip()
        if txt:
            parts.append(txt)
    # tabelas também
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def iter_xlsx_rows(ws, max_rows: int = 4000, max_cols: int = 40) -> Iterable[str]:
    # lê “como texto” com limites de segurança
    row_count = 0
    for row in ws.iter_rows(values_only=True):
        row_count += 1
        if row_count > max_rows:
            yield f"[TRUNCATED: more than {max_rows} rows]"
            break
        # corta colunas
        values = row[:max_cols]
        # filtra vazios
        cleaned = []
        for v in values:
            if v is None:
                continue
            s = str(v).strip()
            if s:
                cleaned.append(s)
        if cleaned:
            yield " | ".join(cleaned)


def read_xlsx(p: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(filename=str(p), read_only=True, data_only=True)
    parts: list[str] = []
    for name in wb.sheetnames:
        ws = wb[name]
        parts.append(f"[SHEET: {name}]")
        for line in iter_xlsx_rows(ws):
            parts.append(line)
        parts.append("")  # separador
    return "\n".join(parts)


def read_any(p: Path) -> str:
    ext = p.suffix.lower()
    if ext in {".txt", ".md"}:
        return read_txt_md(p)
    if ext == ".pdf":
        return read_pdf(p)
    if ext == ".docx":
        return read_docx(p)
    if ext == ".xlsx":
        return read_xlsx(p)
    return ""


# ====== Main ======
def main():
    if not INBOX.exists():
        raise SystemExit(f"Pasta de currículos não encontrada: {INBOX}")

    files = [p for p in INBOX.rglob("*") if p.is_file() and p.suffix.lower() in EXTS]
    print(f"Indexando {len(files)} arquivos em {INBOX} ...")

    for p in files:
        try:
            text = read_any(p)
            text = clean_text(text)
            if not text:
                print(f"SKIP (sem texto): {p.name}")
                continue

            doc_id = file_hash(p)
            delete_by_doc(doc_id)

            chunks = chunk_text(text)
            if not chunks:
                print(f"SKIP (sem chunks): {p.name}")
                continue

            ids = [f"{doc_id}:{i}" for i in range(len(chunks))]
            metas = [
                {
                    "path": str(p),
                    "filename": p.name,
                    "doc_id": doc_id,
                    "chunk": i,
                    "ext": p.suffix.lower(),
                    "folder": str(p.parent),
                    "size_kb": round(p.stat().st_size / 1024, 2),
                }
                for i in range(len(chunks))
            ]

            embs = embed_texts(chunks)
            upsert(ids=ids, docs=chunks, metas=metas, embeddings=embs)

            print(f"OK: {p.name} ({p.suffix.lower()}, {len(chunks)} chunks)")

        except Exception as e:
            print(f"FAIL: {p} -> {e}")

    print("Indexação concluída.")


if __name__ == "__main__":
    main()
